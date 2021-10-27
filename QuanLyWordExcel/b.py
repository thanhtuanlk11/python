from docx import Document 

document = Document()

# số 0 là size chữ
document.add_heading('Giấy mời họp phụ huynh', 0)

p = document.add_paragraph('Học sinh: ')
p.add_run("Nam").bold = True  
p.add_run(" - Lớp 10").italic = True

p = document.add_paragraph('Thời gian: ')
p.add_run("9AM ngày 25/10/2001").bold = True

p = document.add_paragraph('Địa điểm: ')
p.add_run("Phòng a27.1 trường đại học Đà Lạt").bold = True 
document.add_heading('Nội dung',level = 1)

document.add_paragraph('Tổng kết năm học 2021 vừa qua và nhận xét quá trình học cửa từng học sinh')
document.add_paragraph('Lộ trình năm học 2022',style='List Bullet')
document.add_paragraph('Trao thưởng những học sinh xuất sắc', style='List Bullet')
document.add_paragraph('Giải thích các khoảng thu', style='List Bullet')

document.add_heading('Các khoản thu :',level=1)
document.add_paragraph('Học phí : 6 triệu VNĐ', style='List Number')
document.add_paragraph('Đồng phục : 3 triệu VNĐ', style='List Number')
document.add_paragraph('Tham quan : 2 triệu VNĐ', style='List Number')
document.add_heading('Tổng thu : 11 triệu VNĐ', level = 2)


document.save("text.docx") 